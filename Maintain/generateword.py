#-*-coding:utf-8-*-
from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, Inches
import analysis

doc = Document()
doc.styles['Normal'].font.name = u'宋体'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

head = doc.add_paragraph()
head.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = head.add_run('硬件故障分析统计')
run.font.size = Pt(25)


para1 = "本报告主要针对2019年4月22日至2019年4月28日期间发生的服务器硬件故障统计分析。本周共发生服务器硬件故障%s例，其中维修更换配件的硬件故障%s例。涉及到Dell、安擎、联泰、曙光、Amax、浪潮6个品牌，主要故障类型包括Raid、IPMI、GPU、Dmesg、IB、宕机、其他共7个大类。具体统计如下：" %(analysis.error_total_hardwware,analysis.maintain_total[0])
paragraph = doc.add_paragraph(para1)
paragraph.paragraph_format.first_line_indent = Cm(0.74)

para2 = "Dell硬件故障数量：%s\namax硬件故障数量：%s\n安擎硬件故障数量：%s\n联泰硬件故障数量：%s\n曙光硬件故障数量：%s\n浪潮硬件故障数量：%s\n所有硬件故障数量：%s\n"%(analysis.error_dell_hardware,analysis.error_amax_hardware,analysis.error_anqing_hardware,analysis.error_liantai_hardware,analysis.error_shuguang_hardware,analysis.error_langchao_hardware,analysis.error_total_hardwware)
paragraph = doc.add_paragraph(para2)
doc.add_picture('./硬件故障数量.png',width=Inches(4))

#将items数据填到一个表格中
items = (
    ('Dell',analysis.error_dell_Raid,analysis.error_dell_Ipmi,analysis.error_dell_GPU,analysis.error_dell_Dmesg, analysis.error_dell_IB, analysis.error_dell_Dump,analysis.error_dell_Other[0],analysis.error_dell_hardware),
    ('Amax',analysis.error_amax_Raid,analysis.error_amax_Ipmi, analysis.error_amax_GPU,analysis.error_amax_Dmesg, analysis.error_amax_IB, analysis.error_amax_Dump,analysis.error_amax_Other[0],analysis.error_amax_hardware),
    ('安擎',analysis.error_anqing_Raid,analysis.error_anqing_Ipmi, analysis.error_anqing_GPU,analysis.error_anqing_Dmesg, analysis.error_anqing_IB, analysis.error_anqing_Dump,analysis.error_anqing_Other[0],analysis.error_anqing_hardware),
    ('联泰',analysis.error_liantai_Raid,analysis.error_liantai_Ipmi, analysis.error_liantai_GPU,analysis.error_liantai_Dmesg, analysis.error_liantai_IB, analysis.error_liantai_Dump,analysis.error_liantai_Other[0],analysis.error_liantai_hardware),
    ('曙光',analysis.error_shuguang_Raid,analysis.error_shuguang_Ipmi, analysis.error_shuguang_GPU,analysis.error_shuguang_Dmesg, analysis.error_shuguang_IB, analysis.error_shuguang_Dump,analysis.error_shuguang_Other[0],analysis.error_shuguang_hardware),
    ('浪潮',analysis.error_langchao_Raid,analysis.error_langchao_Ipmi, analysis.error_langchao_GPU,analysis.error_langchao_Dmesg, analysis.error_langchao_IB, analysis.error_langchao_Dump,analysis.error_langchao_Other[0],analysis.error_langchao_hardware)
)
titles = ['品牌','Raid','IPMI','GPU','Dmesg','IB','宕机','其他','硬件故障合计']
table = doc.add_table(rows=1,cols=9,style='Medium Grid 1 Accent 1')

for i in range(len(titles)):
    table.rows[0].cells[i].text = titles[i]
for item in items:
    cells = table.add_row().cells
    cells[0].text = item[0]
    cells[1].text = str(item[1])        #注意;不转换为int无法插入，格式不匹配
    cells[2].text = str(item[2])
    cells[3].text = str(item[3])
    cells[4].text = str(item[4])
    cells[5].text = str(item[5])
    cells[6].text = str(item[6])
    cells[7].text = str(item[7])
    cells[8].text = str(item[8])

doc.add_picture('./故障明细对比.png',width=Inches(6))


para3 = "Dell硬件维修故障数量：%s\namax硬件维修故障数量：%s\n安擎维修硬件故障数量：%s\n联泰硬件维修故障数量：%s\n曙光硬件维修故障数量：%s\n浪潮维修硬件故障数量：%s\n所有硬件维修故障数量：%s\n"%(analysis.maintain_dell[0],analysis.maintain_amax[0],analysis.maintain_anqing[0],analysis.maintain_liantai[0],analysis.maintain_shuguang[0],analysis.maintain_langchao[0],analysis.maintain_total[0])
paragraph = doc.add_paragraph(para3)
doc.add_picture('./硬件维修故障数量.png',width=Inches(4))

para4 = "Dell硬件维修故障数量占Dell服务器数量百分比：%s\namax硬件维修故障数量占amax服务器数量百分比：%s\n安擎硬件维修故障数量占安擎服务器数量百分比：%s\n联泰硬件维修故障数量占联泰服务器数量百分比：%s\n曙光硬件维修故障数量占曙光服务器数量百分比：%s\n浪潮硬件维修故障数量占浪潮服务器数量百分比：%s\n所有服务器硬件维修故障数量占服务器数量百分比：%s\n"%(analysis.maintain_percent_dell,analysis.maintain_percent_amax,analysis.maintain_percent_anqing,analysis.maintain_percent_liantai,analysis.maintain_percent_shuguang,analysis.maintain_percent_langchao,analysis.maintain_percent_total)
paragraph = doc.add_paragraph(para4)


#将items数据填到一个表格中
items = (
    ('Dell',analysis.error_dell_Raid_Raidsas[0], analysis.error_dell_Raid_Disk[0], analysis.error_dell_Ipmi_PsError[0],
         analysis.error_dell_Ipmi_PsLoose[0], analysis.error_dell_Ipmi_FanError[0], analysis.error_dell_Ipmi_FanLow[0],
         analysis.error_dell_Ipmi_PCIE[0], analysis.error_dell_Ipmi_Temp[0], analysis.error_dell_Ipmi_Chassis[0],
         analysis.error_dell_GPU_GpuLost[0], analysis.error_dell_GPU_GpuError[0], analysis.error_dell_GPU_GpuBurn[0],
         analysis.error_dell_Dmesg_Mem[0], analysis.error_dell_IB_Cable[0], analysis.error_dell_IB_Card[0],
         analysis.error_dell_ICMP_Dump[0]),
    ('Amax',analysis.error_amax_Raid_Raidsas[0], analysis.error_amax_Raid_Disk[0], analysis.error_amax_Ipmi_PsError[0],
         analysis.error_amax_Ipmi_PsLoose[0], analysis.error_amax_Ipmi_FanError[0], analysis.error_amax_Ipmi_FanLow[0],
         analysis.error_amax_Ipmi_PCIE[0], analysis.error_amax_Ipmi_Temp[0], analysis.error_amax_Ipmi_Chassis[0],
         analysis.error_amax_GPU_GpuLost[0], analysis.error_amax_GPU_GpuError[0], analysis.error_amax_GPU_GpuBurn[0],
         analysis.error_amax_Dmesg_Mem[0], analysis.error_amax_IB_Cable[0], analysis.error_amax_IB_Card[0],
         analysis.error_amax_ICMP_Dump[0]),
    ('安擎',analysis.error_anqing_Raid_Raidsas[0], analysis.error_anqing_Raid_Disk[0], analysis.error_anqing_Ipmi_PsError[0],
         analysis.error_anqing_Ipmi_PsLoose[0], analysis.error_anqing_Ipmi_FanError[0], analysis.error_anqing_Ipmi_FanLow[0],
         analysis.error_anqing_Ipmi_PCIE[0], analysis.error_anqing_Ipmi_Temp[0], analysis.error_anqing_Ipmi_Chassis[0],
         analysis.error_anqing_GPU_GpuLost[0], analysis.error_anqing_GPU_GpuError[0], analysis.error_anqing_GPU_GpuBurn[0],
         analysis.error_anqing_Dmesg_Mem[0], analysis.error_anqing_IB_Cable[0], analysis.error_anqing_IB_Card[0],
         analysis.error_anqing_ICMP_Dump[0]),
    ('联泰',analysis.error_liantai_Raid_Raidsas[0], analysis.error_liantai_Raid_Disk[0], analysis.error_liantai_Ipmi_PsError[0],
         analysis.error_liantai_Ipmi_PsLoose[0], analysis.error_liantai_Ipmi_FanError[0], analysis.error_liantai_Ipmi_FanLow[0],
         analysis.error_liantai_Ipmi_PCIE[0], analysis.error_liantai_Ipmi_Temp[0], analysis.error_liantai_Ipmi_Chassis[0],
         analysis.error_liantai_GPU_GpuLost[0], analysis.error_liantai_GPU_GpuError[0], analysis.error_liantai_GPU_GpuBurn[0],
         analysis.error_liantai_Dmesg_Mem[0], analysis.error_liantai_IB_Cable[0], analysis.error_liantai_IB_Card[0],
         analysis.error_liantai_ICMP_Dump[0]),
    ('曙光',analysis.error_shuguang_Raid_Raidsas[0], analysis.error_shuguang_Raid_Disk[0], analysis.error_shuguang_Ipmi_PsError[0],
         analysis.error_shuguang_Ipmi_PsLoose[0], analysis.error_shuguang_Ipmi_FanError[0], analysis.error_shuguang_Ipmi_FanLow[0],
         analysis.error_shuguang_Ipmi_PCIE[0], analysis.error_shuguang_Ipmi_Temp[0], analysis.error_shuguang_Ipmi_Chassis[0],
         analysis.error_shuguang_GPU_GpuLost[0], analysis.error_shuguang_GPU_GpuError[0], analysis.error_shuguang_GPU_GpuBurn[0],
         analysis.error_shuguang_Dmesg_Mem[0], analysis.error_shuguang_IB_Cable[0], analysis.error_shuguang_IB_Card[0],
         analysis.error_shuguang_ICMP_Dump[0]),
    ('浪潮',analysis.error_langchao_Raid_Raidsas[0],analysis.error_langchao_Raid_Disk[0],analysis.error_langchao_Ipmi_PsError[0],analysis.error_langchao_Ipmi_PsLoose[0],analysis.error_langchao_Ipmi_FanError[0],analysis.error_langchao_Ipmi_FanLow[0],analysis.error_langchao_Ipmi_PCIE[0],analysis.error_langchao_Ipmi_Temp[0],analysis.error_langchao_Ipmi_Chassis[0],analysis.error_langchao_GPU_GpuLost[0],analysis.error_langchao_GPU_GpuError[0],analysis.error_langchao_GPU_GpuBurn[0],analysis.error_langchao_Dmesg_Mem[0],analysis.error_langchao_IB_Cable[0],analysis.error_langchao_IB_Card[0],analysis.error_langchao_ICMP_Dump[0])
)
titles = ["品牌","Raid卡","硬盘故障","电源报错","电源松动","风扇故障","风扇降速","PCIE中断","温度","机箱入侵","GPU掉卡","GPU报错","烤机失败","内存报错","网线故障","IB卡故障","宕机"]
table = doc.add_table(rows=1,cols=17,style='Medium Grid 1 Accent 1')

for i in range(len(titles)):
    table.rows[0].cells[i].text = titles[i]
for item in items:
    cells = table.add_row().cells
    cells[0].text = item[0]
    cells[1].text = str(item[1])        #注意;不转换为int无法插入，格式不匹配
    cells[2].text = str(item[2])
    cells[3].text = str(item[3])
    cells[4].text = str(item[4])
    cells[5].text = str(item[5])
    cells[6].text = str(item[6])
    cells[7].text = str(item[7])
    cells[8].text = str(item[8])
    cells[9].text = str(item[9])
    cells[10].text = str(item[10])
    cells[11].text = str(item[11])
    cells[12].text = str(item[12])
    cells[13].text = str(item[13])
    cells[14].text = str(item[14])
    cells[15].text = str(item[15])
    cells[16].text = str(item[16])


doc.save('硬件故障统计分析.docx')